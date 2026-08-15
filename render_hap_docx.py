"""
HAP Word Rendering
==================
Converts a verified HAP working draft into the reader-facing Word document.

The working draft (<slug>_hap_draft.md) keeps machine-checkable markup that
verify_hap_numbers.py depends on. This script produces the presentation copy:

  [VERIFY: claim]  ->  claim*            (asterisk = "verify before publishing")
  [BLANK: what]    ->  To be completed by city officials — what.
  [MAP: ...]       ->  italic map-insertion placeholder (if any remain)

plus a legend for the asterisk at the end of the document, then runs pandoc.

Usage:
    python3 render_hap_docx.py --city Agra
Reads  outputs/<slug>_hap_draft.md
Writes outputs/<slug>_hap_render.md and outputs/<City>_HAP_Draft.docx
"""

import argparse
import re
import subprocess
import sys

LEGEND = ('\\* *Statement drawn from public sources or model knowledge; '
          'to be verified by city officials before publication.*')


def style_tables(docx_path):
    """Add visible gridlines and a shaded header row to every table.

    Pandoc's default reference doc renders tables nearly borderless, which
    reviewers flagged as hard to read. Injects tblBorders + header shading
    directly so no custom reference.docx is needed.
    """
    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError:
        print('WARNING: python-docx not installed — tables left unstyled '
              '(pip3 install python-docx)')
        return

    doc = Document(docx_path)
    for table in doc.tables:
        tbl_pr = table._tbl.tblPr
        for old in tbl_pr.findall(qn('w:tblBorders')):
            tbl_pr.remove(old)
        borders = OxmlElement('w:tblBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '4')          # 0.5 pt
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), '808080')
            borders.append(el)
        tbl_pr.append(borders)
        if table.rows:
            for cell in table.rows[0].cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:fill'), 'E8E8E8')
                tc_pr.append(shd)
    doc.save(docx_path)
    print(f'styled {len(doc.tables)} tables (gridlines + header shading)')


LIST_MARKER = re.compile(r'^(\s*)([-*+]|\d+[.)])\s')


def ensure_blank_before_lists(text):
    """Pandoc only starts a new list block after a blank line; a lead-in
    line glued directly to its bullets (no blank line between) gets merged
    into one run-on paragraph in the .docx — bullets silently vanish, only
    visible by inspecting the actual Word file, not the markdown source.
    Insert the missing blank line wherever that pattern occurs, so this
    can't happen regardless of what the draft (LLM-written or hand-edited)
    does.
    """
    lines = text.split('\n')
    out = []
    for i, line in enumerate(lines):
        out.append(line)
        if (i + 1 < len(lines) and line.strip() and not LIST_MARKER.match(line)
                and LIST_MARKER.match(lines[i + 1])):
            out.append('')
    return '\n'.join(out)


def render(text):
    text = ensure_blank_before_lists(text)

    # Standard Early Warning System block — identical across every city, so
    # it is inserted here rather than left to the LLM to redraft each time.
    # Its own [BLANK: ...] tokens flow into the normal conversion below.
    if '[EARLY_WARNING_TEMPLATE]' in text:
        template = open('hap_early_warning_template.md').read()
        text = text.replace('[EARLY_WARNING_TEMPLATE]', template.strip())

    n_verify = 0

    def _verify(m):
        nonlocal n_verify
        n_verify += 1
        return m.group(1).strip().rstrip('.') + '\\*'

    def _blank(m):
        what = m.group(1).strip().rstrip('.')
        return f'*To be completed by city officials — {what}.*'

    def _map(m):
        return f'*Map to be inserted: {m.group(1).strip()}*'

    text = re.sub(r'\[VERIFY:\s*(.*?)\]', _verify, text, flags=re.DOTALL)
    text = re.sub(r'\[BLANK:\s*(.*?)\]', _blank, text, flags=re.DOTALL)
    text = re.sub(r'\[MAP:\s*(.*?)\]', _map, text, flags=re.DOTALL)

    if n_verify:
        text = text.rstrip() + '\n\n---\n\n' + LEGEND + '\n'
    return text, n_verify


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--city', default='Agra')
    args = ap.parse_args()
    city, slug = args.city, args.city.lower()

    draft = open(f'outputs/{slug}_hap_draft.md').read()
    rendered, n_verify = render(draft)
    render_path = f'outputs/{slug}_hap_render.md'
    open(render_path, 'w').write(rendered)
    print(f'wrote {render_path} ({n_verify} verify-marked statements -> "*")')

    docx = f'outputs/{city}_HAP_Draft.docx'
    r = subprocess.run(['pandoc', render_path, '-o', docx,
                        '--resource-path=outputs',
                        '--metadata', f'title=Heat Action Plan for {city} — DRAFT'])
    if r.returncode:
        sys.exit('pandoc conversion failed (is pandoc installed?).')
    style_tables(docx)
    print(f'wrote {docx}')


if __name__ == '__main__':
    main()
